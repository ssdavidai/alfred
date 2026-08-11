"""file_extraction — activities for the FileExtractionWorkflow.

Background
----------

Issue #114 Lane B closes the load-bearing UX promise of §14 step 3:
after a principal drops a file on /files, Alfred has 30 seconds to
read it and surface a one-paragraph summary + an "Alfred read it"
badge. The audit at /tmp/orchestrator-114-coverage.md flagged this as
the only blocker for steps #3–#5 of the cold-start test.

This module is the alfred-learn glue. ctrl-api owns the blob volume
and the SQL writes; we read the bytes back over HTTP, run the
per-mime extractor, ask the workers gateway for a one-paragraph
summary, and PATCH the result back via
``/api/v1/files/:id/extraction``. Same discipline ``files_cold_archive``
already follows.

Three activities
----------------

* ``read_file_metadata(file_id)`` — pulls the canonical row from
  ctrl-api so the workflow knows the mime + path. Cheap GET.
* ``fetch_and_extract_text(file_id, content_type, path)`` — fetches
  the blob bytes and dispatches to the right extractor by mime.
  Returns ``{text, char_count, extractor, truncated_for_summary}``
  on success or raises a classified ``ExtractionError`` with a short
  reason code on failure.
* ``summarise_extracted_text(...)`` — workers gateway one-shot summary.
* ``stamp_extraction_result(file_id, alfred_read_at, summary,
  extraction_error)`` — write-back PATCH.

Per-mime extractor matrix (Lane B scope)
----------------------------------------

| Mime                                                                              | Extractor    |
|-----------------------------------------------------------------------------------|--------------|
| ``text/*``, ``application/json``, ``application/xml``, ``application/x-yaml``     | direct utf-8 |
| ``application/pdf``                                                                | pypdf        |
| ``application/vnd.openxmlformats-officedocument.wordprocessingml.document``        | python-docx  |
| ``application/vnd.openxmlformats-officedocument.spreadsheetml.sheet``              | openpyxl     |
| everything else                                                                    | unsupported  |

OCR (tesseract) + audio (Groq Whisper) are intentionally outside Lane
B's scope — the spec calls for them but they each merit their own PR
(image OCR is a separate apt package; Whisper is a Groq quota line
item). The workflow marks them ``extraction_error="unsupported_mime"``
so the dashboard's subtle-error state can render a "Alfred can't read
this yet" affordance.

Environment
-----------

* ``ALFRED_CTRL_URL`` — the tenant ctrl-api base URL. Defaults to
  ``http://alfred-ctrl:3100``. Same knob ``StateClient`` /
  ``VaultClient`` honour.
* ``AAS_API_KEY`` — the operator bearer token. Required.
* ``FILES_EXTRACTION_MAX_CHARS`` — soft cap on extractor output handed
  to the summariser (default 60_000, ~15 K tokens of input). Bigger
  files are truncated; the summariser is told.
"""
from __future__ import annotations

import io
import logging
import os
from typing import Any

import httpx
from temporalio import activity

logger = logging.getLogger("alfred-learn")


# Same env knob StateClient / VaultClient honour.
_DEFAULT_CTRL_URL = "http://alfred-ctrl:3100"

# Soft cap on extracted-text length handed to the summariser. The
# workers gateway is generous (clerk.py allows 31 K output tokens) but
# the INPUT side competes with every other autonomous job for the same
# gateway; cap at ~60 KB (~15 K tokens) so a 100-page PDF doesn't
# starve the rest of the queue. A truncation notice is appended so the
# summariser doesn't hallucinate completeness.
_MAX_CHARS_FOR_SUMMARY = int(os.environ.get("FILES_EXTRACTION_MAX_CHARS", "60000"))

# Tightish blob fetch — a multi-MB PDF on a fresh tenant downloads in
# <1s over localhost loopback. 30s is the same envelope ctrl-api gives
# itself for the cold-promote sweep.
_BLOB_FETCH_TIMEOUT_S = 30.0

# Mime classes the extractor can handle directly. The workflow uses
# these constants to short-circuit "unsupported" before paying for a
# blob fetch.
_TEXT_MIMES = (
    "text/",
    "application/json",
    "application/xml",
    "application/x-yaml",
    "application/yaml",
)
_PDF_MIME = "application/pdf"
_DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_XLSX_MIME = (
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)


# ── error taxonomy ─────────────────────────────────────────────────────


class ExtractionError(Exception):
    """An expected, recoverable extraction failure.

    The ``code`` attribute is the short reason that lands in
    ``files.extraction_error``; it's what the dashboard's tooltip
    will read. ``str(exc)`` always starts with ``code + ": "`` so a
    Temporal-wrapped error envelope can recover the code via a simple
    head-split even when the original ``code`` attribute didn't
    survive serialisation.
    """

    def __init__(self, code: str, message: str = "") -> None:
        # Always emit ``"<code>: <message>"`` (or just ``"<code>"`` if no
        # message) so the workflow's ``_extraction_error_code`` recovery
        # path can split on the first colon. This is the contract the
        # workflow's wrap-handling relies on — keep it stable.
        wire = f"{code}: {message}" if message else code
        super().__init__(wire)
        self.code = code


# ── helpers ────────────────────────────────────────────────────────────


def _ctrl_base_url() -> str:
    return os.environ.get("ALFRED_CTRL_URL", _DEFAULT_CTRL_URL)


def _auth_headers() -> dict[str, str]:
    key = os.environ.get("AAS_API_KEY", "")
    return {"Authorization": f"Bearer {key}"} if key else {}


def classify_mime(content_type: str | None) -> str:
    """Return a short token describing how the workflow should handle
    this mime. One of ``"text"``, ``"pdf"``, ``"docx"``, ``"xlsx"``,
    ``"unsupported"``.

    Pure: no httpx, no filesystem. Tested directly.
    """
    ct = (content_type or "").strip().lower()
    if not ct:
        return "unsupported"
    # Strip any "; charset=…" tail.
    head = ct.split(";", 1)[0].strip()
    if head == _PDF_MIME:
        return "pdf"
    if head == _DOCX_MIME:
        return "docx"
    if head == _XLSX_MIME:
        return "xlsx"
    for prefix in _TEXT_MIMES:
        if head.startswith(prefix):
            return "text"
    return "unsupported"


def _truncate_for_summary(text: str) -> tuple[str, bool]:
    """Truncate ``text`` to the summariser cap, returning the trimmed
    string + a flag indicating whether truncation happened."""
    if len(text) <= _MAX_CHARS_FOR_SUMMARY:
        return text, False
    head = text[:_MAX_CHARS_FOR_SUMMARY]
    return head, True


# ── extractors ─────────────────────────────────────────────────────────


def extract_text_blob(content_type: str | None, blob: bytes) -> str:
    """Pure extractor dispatch.

    ``blob`` is the file bytes; the function returns the extracted
    plain text or raises ``ExtractionError`` with a short reason code.

    Tested directly against fixture bytes — no httpx, no Temporal.
    """
    kind = classify_mime(content_type)
    if kind == "unsupported":
        raise ExtractionError(
            "unsupported_mime",
            f"no Lane-B extractor for content_type={content_type!r}",
        )
    try:
        if kind == "text":
            return _extract_text(blob, content_type)
        if kind == "pdf":
            return _extract_pdf(blob)
        if kind == "docx":
            return _extract_docx(blob)
        if kind == "xlsx":
            return _extract_xlsx(blob)
    except ExtractionError:
        raise
    except Exception as exc:  # noqa: BLE001 — extractor wrappers are noisy
        raise ExtractionError(
            "extractor_failed",
            f"{kind} extractor raised: {exc}",
        ) from exc
    # Defensive — classify_mime guarantees one of the above.
    raise ExtractionError("extractor_failed", f"unknown extractor kind={kind!r}")


def _extract_text(blob: bytes, content_type: str | None) -> str:
    """UTF-8 decode with a small fallback chain. Stops short of full
    chardet — text/* files in the wild are overwhelmingly UTF-8 today,
    and the cost of guessing wrong on a binary blob is much worse than
    a one-off 'extractor_failed' on a Latin-1 file.
    """
    # Cheap fast path.
    try:
        return blob.decode("utf-8")
    except UnicodeDecodeError:
        pass
    # Tolerate Latin-1 / Windows-1252 by replacing unmappable bytes.
    try:
        return blob.decode("utf-8", errors="replace")
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(
            "extractor_failed",
            f"text decode failed (ct={content_type!r}): {exc}",
        ) from exc


def _extract_pdf(blob: bytes) -> str:
    """pypdf page-by-page text extraction.

    pypdf is pure Python (no native poppler dep) and handles the
    common-case PDF in a few hundred ms per page. Scanned PDFs (no
    embedded text layer) return empty strings per page — we surface
    that as a distinct ``empty_pdf`` error code so the dashboard can
    say "no text layer; OCR needed" rather than "extraction failed".
    """
    try:
        from pypdf import PdfReader  # local import — keeps cold start fast
    except ImportError as exc:  # pragma: no cover — defensive
        raise ExtractionError(
            "extractor_failed",
            f"pypdf import failed: {exc}",
        ) from exc

    try:
        reader = PdfReader(io.BytesIO(blob))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(
            "extractor_failed", f"pypdf could not parse PDF: {exc}",
        ) from exc

    pieces: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            chunk = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001
            logger.warning(
                "file_extraction: pdf page %d extract_text raised: %s", i, exc
            )
            chunk = ""
        if chunk:
            pieces.append(chunk.strip())
    out = "\n\n".join(p for p in pieces if p)
    if not out.strip():
        raise ExtractionError(
            "empty_pdf",
            "PDF has no embedded text layer (scanned image — OCR needed)",
        )
    return out


def _extract_docx(blob: bytes) -> str:
    """python-docx paragraph + table text. Drops images / shapes — the
    summariser cares about prose, not formatting."""
    try:
        from docx import Document  # local import
    except ImportError as exc:  # pragma: no cover — defensive
        raise ExtractionError(
            "extractor_failed",
            f"python-docx import failed: {exc}",
        ) from exc
    try:
        doc = Document(io.BytesIO(blob))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(
            "extractor_failed", f"python-docx could not parse: {exc}",
        ) from exc
    pieces: list[str] = []
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            pieces.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                pieces.append(" | ".join(cells))
    out = "\n".join(pieces)
    if not out.strip():
        raise ExtractionError(
            "empty_docx", "DOCX has no extractable paragraph or table text",
        )
    return out


def _extract_xlsx(blob: bytes) -> str:
    """openpyxl read_only sheet scan. Each sheet contributes a banner
    line + one CSV-ish row per non-empty row. read_only mode keeps RAM
    bounded on multi-MB spreadsheets."""
    try:
        from openpyxl import load_workbook  # local import
    except ImportError as exc:  # pragma: no cover — defensive
        raise ExtractionError(
            "extractor_failed",
            f"openpyxl import failed: {exc}",
        ) from exc
    try:
        wb = load_workbook(io.BytesIO(blob), read_only=True, data_only=True)
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(
            "extractor_failed", f"openpyxl could not parse: {exc}",
        ) from exc
    pieces: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        pieces.append(f"## Sheet: {sheet_name}")
        for row in ws.iter_rows(values_only=True):
            non_empty = [str(v) for v in row if v not in (None, "")]
            if non_empty:
                pieces.append(",".join(non_empty))
    try:
        wb.close()
    except Exception:  # noqa: BLE001
        pass
    # An XLSX with zero non-empty cells is still a degenerate-but-valid
    # input — fail it the same way pdf/docx do.
    body = "\n".join(p for p in pieces if not p.startswith("## Sheet"))
    if not body.strip():
        raise ExtractionError(
            "empty_xlsx", "XLSX has no non-empty cells",
        )
    return "\n".join(pieces)


# ── activity wrappers ──────────────────────────────────────────────────


@activity.defn
async def read_file_metadata(file_id: str) -> dict[str, Any]:
    """Pull one row's metadata from ctrl-api by file id.

    Used by the workflow to learn the mime + path before deciding
    whether to even fetch the blob (unsupported mimes short-circuit).
    """
    # ctrl-api has no GET-by-id surface today — the existing list
    # endpoint already supports pagination but no id filter, so we
    # fetch a tight slice and filter client-side. This keeps Lane B
    # from coupling to a not-yet-shipped /by-id route.
    url = f"{_ctrl_base_url()}/api/v1/files/list"
    async with httpx.AsyncClient(timeout=10.0) as client:
        resp = await client.get(
            url,
            params={"limit": "500"},
            headers=_auth_headers(),
        )
        resp.raise_for_status()
        payload = resp.json()
    items = payload.get("items") or []
    for item in items:
        if str(item.get("id") or "") == file_id:
            return item
    raise ExtractionError("not_found", f"file_id={file_id} not in /list")


@activity.defn
async def fetch_and_extract_text(
    file_id: str, content_type: str | None, path: str,
) -> dict[str, Any]:
    """Fetch the blob + extract text. Returns
    ``{text, char_count, extractor, truncated_for_summary}``.

    ``ExtractionError`` raises propagate to the workflow's stamp step
    — they're classified failures, not Temporal retryable errors.
    """
    kind = classify_mime(content_type)
    if kind == "unsupported":
        raise ExtractionError(
            "unsupported_mime",
            f"Lane-B does not handle content_type={content_type!r}",
        )
    # URL-encode each path segment defensively (the blob route
    # registered as `/blob/*` takes the literal tail). urllib's
    # quote(safe="") gives us per-segment control without rebuilding
    # the absolute URL.
    from urllib.parse import quote

    encoded_path = "/".join(
        quote(seg, safe="") for seg in path.split("/") if seg
    )
    blob_url = f"{_ctrl_base_url()}/api/v1/files/blob/{encoded_path}"
    async with httpx.AsyncClient(timeout=_BLOB_FETCH_TIMEOUT_S) as client:
        resp = await client.get(blob_url, headers=_auth_headers())
        resp.raise_for_status()
        blob = resp.content
    text = extract_text_blob(content_type, blob)
    truncated_text, truncated = _truncate_for_summary(text)
    return {
        "text": truncated_text,
        "char_count": len(text),
        "extractor": kind,
        "truncated_for_summary": truncated,
    }


def _build_summary_prompt(
    *,
    text: str,
    original_filename: str | None,
    content_type: str | None,
    truncated: bool,
) -> str:
    """Render the one-paragraph summary prompt for the workers gateway.

    Kept as a module-level helper so it's unit-testable + so the
    activity body stays linear.
    """
    name_line = (
        f"FILENAME: {original_filename}\n" if original_filename else ""
    )
    truncate_line = (
        "\nNote: the file was truncated to fit the summariser's input "
        "budget; the body below is the first portion of the document.\n"
        if truncated
        else ""
    )
    return f"""You are summarising a file the principal just dropped into their personal file store. Write ONE short paragraph (3-5 sentences, <=600 chars) describing what this file is and what's in it — the kind of thing a butler would whisper while handing over the document.

Rules:
- Plain prose, no markdown, no bullet lists.
- Lead with what the document IS (invoice, contract, screenshot, code listing, ...).
- Then the one or two facts a reader would want next (parties, dates, amounts, totals).
- No filler ("This document contains ..."). No hedging.
- If the body is empty or gibberish, say so plainly.

Return JSON only:
{{"summary": "<the paragraph>"}}

{name_line}CONTENT_TYPE: {content_type or "unknown"}
{truncate_line}
BODY:
{text}
"""


@activity.defn
async def summarise_extracted_text(
    file_id: str,
    extracted_text: str,
    original_filename: str | None,
    content_type: str | None,
    truncated: bool,
) -> str:
    """Ask the workers gateway for the one-paragraph summary.

    Routes through ``clerk._call_clerk`` so the Temporal activity gets
    the heartbeat + per-call billing classification the rest of the
    learn pipeline already enjoys.
    """
    # Lazy import — avoids the clerk module being loaded by tests that
    # exercise the pure extractor only.
    from src.activities.clerk import _call_clerk

    prompt = _build_summary_prompt(
        text=extracted_text,
        original_filename=original_filename,
        content_type=content_type,
        truncated=truncated,
    )
    try:
        result = await _call_clerk(prompt, agent_id=f"file-extract-{file_id}")
    except Exception as exc:  # noqa: BLE001
        # Workers gateway threw — network error, HTTP 5xx, timeout, billing.
        # These are transient; the file should be re-tried when the gateway
        # recovers.  Use a distinct code so operators can tell this apart from
        # a structural failure where the model returned the wrong shape.
        raise ExtractionError(
            "summariser_gateway_error", f"workers gateway raised: {exc}",
        ) from exc
    summary: Any = None
    if isinstance(result, dict):
        summary = result.get("summary")
    elif isinstance(result, str):
        summary = result
    if not isinstance(summary, str) or not summary.strip():
        # Gateway ran successfully but the model returned JSON without a
        # "summary" key (or with an empty value).  This is more likely a
        # prompt-shape issue than a transient availability problem — kept
        # separate so it can be investigated independently.
        raise ExtractionError(
            "summariser_no_output", "workers gateway returned no summary text",
        )
    # Hard cap (the ctrl-api PATCH route caps at 4 KiB; we trim earlier
    # so a chatty model never breaches it).
    summary = summary.strip()
    if len(summary) > 4096:
        summary = summary[:4093] + "..."
    return summary


@activity.defn
async def stamp_extraction_result(
    file_id: str,
    alfred_read_at: int | None,
    summary: str | None,
    extraction_error: str | None,
) -> dict[str, Any]:
    """PATCH the result back to ctrl-api.

    Always emits all three columns — the route accepts partial
    patches but a fresh extraction always wants to clear any prior
    error and stamp the new state in one shot.
    """
    url = f"{_ctrl_base_url()}/api/v1/files/{file_id}/extraction"
    body = {
        "alfred_read_at": alfred_read_at,
        "summary": summary,
        "extraction_error": extraction_error,
    }
    async with httpx.AsyncClient(timeout=15.0) as client:
        resp = await client.patch(url, json=body, headers=_auth_headers())
        resp.raise_for_status()
        return resp.json()
